"""

Convert PRE-BEP.md to Word document using template

"""
import os

from pathlib import Path

from docx import Document

from docx.shared import Pt, Inches, RGBColor

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml.ns import qn

from docx.oxml import OxmlElement

from docx.oxml.shared import qn as shared_qn
import re


# Paths

BASE_DIR = Path(__file__).parent.parent

INPUT_PREBEP_MD = BASE_DIR / "PRE-BEP.md"
TEMPLATE_PREBEP_DOCX = BASE_DIR / "input" / "word" / "PRE_BEP_template.docx"
OUTPUT_PREBEP_DOCX = BASE_DIR / "output" / "BIM protokol_Příloha 03_PRE-BEP (šablona).docx"

INPUT_EIR_MD = BASE_DIR / "EIR.md"
TEMPLATE_EIR_DOCX = BASE_DIR / "input" / "word" / "EIR_template.docx"
OUTPUT_EIR_DOCX = BASE_DIR / "output" / "BIM protokol_Příloha 02_EIR.docx"

INPUT_PROTOKOL_MD = BASE_DIR / "BIM protokol.md"
TEMPLATE_PROTOKOL_DOCX = BASE_DIR / "input" / "word" / "PROTOKOL_template.docx"
OUTPUT_PROTOKOL_DOCX = BASE_DIR / "output" / "BIM protokol.docx"

INPUT_PRILOHA_MD = BASE_DIR / "Priloha-B.md"
TEMPLATE_PRILOHA_DOCX = BASE_DIR / "input" / "word" / "EIR_template.docx"
OUTPUT_PRILOHA_DOCX = BASE_DIR / "output" / "EIR_Priloha_B_Geometrická podrobnost prvků.docx"

	

def read_markdown_file(filepath):

	"""Read markdown file and return content"""

	with open(filepath, 'r', encoding='utf-8') as f:

		return f.read()



def parse_table(lines, start_idx):

	"""Parse markdown table starting at given index"""

	table_lines = []

	idx = start_idx
	

	while idx < len(lines) and '|' in lines[idx]:

		table_lines.append(lines[idx])

		idx += 1
	

	if len(table_lines) < 2:

		return None, start_idx
	

	# Parse header

	header_line = table_lines[0]

	headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
	

	# Skip separator line, parse all data rows (including empty ones)

	rows = []

	for line in table_lines[2:]:

		cells = [cell.strip() for cell in line.split('|')[1:-1]]

		# Add all rows, even if empty (to preserve table structure)

		rows.append(cells)
	

	return {'headers': headers, 'rows': rows}, idx



def add_table_to_doc(doc, table_data):

	"""Add table to Word document"""

	if not table_data or not table_data['rows']:

		return
	

	num_cols = len(table_data['headers'])

	num_rows = len(table_data['rows']) + 1  # +1 for header
	

	table = doc.add_table(rows=num_rows, cols=num_cols)
	

	# Set dark grey borders for the table

	from docx.oxml import parse_xml

	tbl = table._element

	tblPr = tbl.find(qn('w:tblPr'))

	if tblPr is None:

		tblPr = OxmlElement('w:tblPr')

		tbl.insert(0, tblPr)
	

	# Create table borders with dark grey color

	tblBorders = parse_xml(

		'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'

		'<w:top w:val="single" w:sz="4" w:space="0" w:color="808080"/>'

		'<w:left w:val="single" w:sz="4" w:space="0" w:color="808080"/>'

		'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="808080"/>'

		'<w:right w:val="single" w:sz="4" w:space="0" w:color="808080"/>'

		'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="808080"/>'

		'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="808080"/>'

		'</w:tblBorders>'
	)

	tblPr.append(tblBorders)
	

	# Add headers with Tabulka - záhlaví style

	header_cells = table.rows[0].cells

	for idx, header in enumerate(table_data['headers']):

		cleaned_header = clean_markdown_formatting(header)

		header_cells[idx].text = cleaned_header

		# Apply Tabulka - záhlaví style if available

		for paragraph in header_cells[idx].paragraphs:

			try:

				paragraph.style = 'Tabulka - záhlaví'

			except KeyError:

				# Style not available, make bold as fallback

				for run in paragraph.runs:

					run.bold = True
	

	# Add data rows with Tabulka style

	for row_idx, row_data in enumerate(table_data['rows']):

		cells = table.rows[row_idx + 1].cells

		for col_idx, cell_data in enumerate(row_data):

			if col_idx < num_cols:

				# Clean markdown and handle line breaks

				cleaned_data = clean_markdown_formatting(cell_data)

				cells[col_idx].text = cleaned_data.replace('<br>', '\n')

				# Apply Tabulka style if available

				for paragraph in cells[col_idx].paragraphs:

					try:

						paragraph.style = 'Tabulka'

					except KeyError:

						pass  # Use default style



def add_table_of_contents(paragraph, levels=2):

	"""Add a table of contents field to the given paragraph"""

	run = paragraph.add_run()
	

	# Create the TOC field code

	fldChar1 = OxmlElement('w:fldChar')

	fldChar1.set(qn('w:fldCharType'), 'begin')
	

	instrText = OxmlElement('w:instrText')

	instrText.set(qn('xml:space'), 'preserve')

	instrText.text = f'TOC \\o "1-{levels}" \\h \\z \\u'
	

	fldChar2 = OxmlElement('w:fldChar')

	fldChar2.set(qn('w:fldCharType'), 'separate')
	

	fldChar3 = OxmlElement('w:fldChar')

	fldChar3.set(qn('w:fldCharType'), 'end')
	

	# Add the field code elements to the run

	run._r.append(fldChar1)

	run._r.append(instrText)

	run._r.append(fldChar2)

	run._r.append(fldChar3)



def clean_markdown_formatting(text):

	"""Remove markdown formatting from text"""

	# Remove bold markers (strip them, don't convert to bold)

	text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)

	# Remove italic markers

	text = re.sub(r'\*(.+?)\*', r'\1', text)

	# Remove code markers

	text = re.sub(r'`(.+?)`', r'\1', text)

	# Remove links but keep text (will be handled separately for internal links)

	text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)

	return text



def generate_bookmark_name(text):

	"""Generate a valid bookmark name from heading text"""

	# Remove special characters and replace spaces with underscores

	bookmark = re.sub(r'[^\w\s-]', '', text)

	bookmark = re.sub(r'\s+', '_', bookmark)

	# Ensure it starts with a letter and is max 40 chars

	if not bookmark[0].isalpha():

		bookmark = 'BM_' + bookmark

	return bookmark[:40]



def add_bookmark_to_paragraph(paragraph, bookmark_name):

	"""Add a bookmark to a paragraph"""

	from docx.oxml import parse_xml
	

	p_element = paragraph._element
	

	# Create bookmark start

	bookmark_start = parse_xml(

		f'<w:bookmarkStart w:id="0" w:name="{bookmark_name}" '

		'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
	)
	

	# Create bookmark end

	bookmark_end = parse_xml(

		'<w:bookmarkEnd w:id="0" '

		'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
	)
	

	# Insert bookmark around the paragraph content

	p_element.insert(0, bookmark_start)

	p_element.append(bookmark_end)



def process_text_with_links(paragraph, text, heading_bookmarks):

	"""Process text and convert internal markdown links to Word hyperlinks"""

	# Find all internal links in the format [text](#anchor)

	link_pattern = r'\[([^\]]+)\]\(#([^\)]+)\)'
	

	parts = []

	last_end = 0
	

	for match in re.finditer(link_pattern, text):

		# Add text before the link

		if match.start() > last_end:

			parts.append(('text', text[last_end:match.start()]))
		

		# Add the link

		link_text = match.group(1)

		anchor = match.group(2)

		parts.append(('link', link_text, anchor))
		

		last_end = match.end()
	

	# Add remaining text

	if last_end < len(text):

		parts.append(('text', text[last_end:]))
	

	# If no links found, just add the text

	if not parts:

		paragraph.add_run(text)

		return
	

	# Build the paragraph with links

	for part in parts:

		if part[0] == 'text':

			paragraph.add_run(part[1])

		elif part[0] == 'link':

			link_text = part[1]

			anchor = part[2]
			

			# Check if we have a bookmark for this anchor

			if anchor in heading_bookmarks:

				bookmark_name = heading_bookmarks[anchor]

				add_hyperlink(paragraph, link_text, bookmark_name)

			else:

				# If no bookmark found, just add as regular text

				paragraph.add_run(link_text)



def add_hyperlink(paragraph, text, bookmark_name):

	"""Add a hyperlink to a bookmark within the paragraph"""

	from docx.oxml import parse_xml

	from docx.oxml.ns import qn
	

	# Get the paragraph element

	p = paragraph._element
	

	# Create hyperlink element with the run inside it

	# Apply blue color and underline directly to ensure hyperlink appearance

	hyperlink_xml = (

		f'<w:hyperlink w:anchor="{bookmark_name}" '

		'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'

		'<w:r>'

		'<w:rPr>'

		'<w:color w:val="156082"/>'

		'<w:u w:val="single"/>'

		'</w:rPr>'

		f'<w:t>{text}</w:t>'

		'</w:r>'

		'</w:hyperlink>'
	)
	

	hyperlink = parse_xml(hyperlink_xml)

	p.append(hyperlink)



def remove_header_numbering(text):

	"""Remove numbering from header text (e.g., '1.', '1.1.', '1.1.1.', '7.9.1	Text')"""

	# Remove patterns like "1.", "1.1.", "1.1.1.", "1.1.1.1." at the start

	# Also handles numbers without trailing period like "7.9.1	Text"

	text = re.sub(r'^\d+(\.\d+)*\.?\s+', '', text)

	return text



def apply_list_numbering(paragraph, num_id=10, level=0):

	"""Apply numbering to a paragraph. num_id: numbering definition (10=simple list, high number to avoid header conflicts), level: 0-based indent level"""

	p_element = paragraph._element

	pPr = p_element.get_or_add_pPr()

	

	# Remove any existing numPr to avoid conflicts

	existing_numPr = pPr.find(qn('w:numPr'))

	if existing_numPr is not None:

		pPr.remove(existing_numPr)

	

	# Create new numPr element

	numPr = OxmlElement('w:numPr')

	pPr.append(numPr)

	

	# Set numId (numbering definition ID - use high number to avoid conflict with header numbering)

	numId = OxmlElement('w:numId')

	numId.set(qn('w:val'), str(num_id))

	numPr.append(numId)

	

	# Set ilvl (indentation level, 0-based)

	ilvl = OxmlElement('w:ilvl')

	ilvl.set(qn('w:val'), str(level))

	numPr.append(ilvl)

	numPr.append(ilvl)



def extract_metadata_value(line):

	"""Extract value from metadata line like **Projekt**: `[Název projektu]`"""

	# Match pattern: **Label**: `value` or **Label**: value

	match = re.match(r'\*\*(.+?)\*\*:\s*`(.+?)`', line)

	if match:

		return match.group(1), match.group(2)

	match = re.match(r'\*\*(.+?)\*\*:\s*(.+)', line)

	if match:

		return match.group(1), match.group(2).strip()

	return None, None



def set_company_property(doc, company_name):

	"""Set the Company property using category"""

	doc.core_properties.category = company_name



def convert_markdown_to_word(md_file, template_file, output_file):

	"""Convert markdown file to Word document using template"""
	

	# Read markdown content

	print(f"Reading markdown file: {md_file}")

	content = read_markdown_file(md_file)

	lines = content.split('\n')
	

	# Load template

	print(f"Loading template: {template_file}")

	doc = Document(template_file)
	

	# Extract document title (first line with # hashtag) and metadata from the beginning

	doc_title = None

	metadata = {}

	metadata_end_idx = 0
	

	for i, line in enumerate(lines):

		if line.startswith('# ') and doc_title is None:

			# First header with hashtag - extract title

			doc_title = clean_markdown_formatting(line[2:])

			doc_title = remove_header_numbering(doc_title)

		elif line.startswith('**') and ':' in line:

			label, value = extract_metadata_value(line)

			if label:

				metadata[label] = value

				metadata_end_idx = i + 1

		elif line.startswith('# ') and doc_title is not None:

			# Second header - stop metadata extraction

			break

		elif line.strip() == '':

			continue

		else:

			# Stop at first non-metadata, non-title line

			if doc_title and metadata:

				break
	

	# Set document properties

	print(f"Setting document properties - Title: {doc_title}, Metadata: {metadata}")

	if doc_title:

		doc.core_properties.title = doc_title  # Name property

	if 'Projekt' in metadata:

		doc.core_properties.subject = metadata['Projekt']  # Subject property

	if 'Objednatel' in metadata:

		set_company_property(doc, metadata['Objednatel'])  # Company property

	if 'Verze' in metadata:

		doc.core_properties.keywords = metadata['Verze']  # Keywords property
	

	# First pass: build a mapping of markdown anchors to bookmark names

	# This maps anchor IDs (e.g., "52-složení-realizačního-týmu") to bookmark names

	heading_bookmarks = {}

	skip_first_header_pass = True
	

	for line in lines:

		if line.startswith('# ') or line.startswith('## ') or line.startswith('### ') or line.startswith('#### '):

			# Skip first header (document title)

			if skip_first_header_pass:

				skip_first_header_pass = False

				continue
			

			# Extract heading text

			level = line.count('#', 0, 4)

			heading_text_raw = clean_markdown_formatting(line.lstrip('#').strip())
			

			# Generate anchor ID from raw heading (with numbers) - same format as markdown

			anchor_id = heading_text_raw.lower()

			anchor_id = re.sub(r'[^\w\s-]', '', anchor_id)  # Remove special chars (including dots)

			anchor_id = re.sub(r'\s+', '-', anchor_id)  # Replace spaces with dashes

			anchor_id = re.sub(r'-+', '-', anchor_id)  # Replace multiple dashes with single

			anchor_id = anchor_id.strip('-')  # Remove leading/trailing dashes
			

			# Generate bookmark name from cleaned heading (without numbers)

			heading_text = remove_header_numbering(heading_text_raw)

			bookmark_name = generate_bookmark_name(heading_text)
			

			# Store mapping

			heading_bookmarks[anchor_id] = bookmark_name

			print(f"Mapped anchor '{anchor_id}' to bookmark '{bookmark_name}'")
	

	# Parse markdown and add to document

	print("Converting markdown to Word...")
	

	idx = 0

	in_metadata_section = True

	skip_toc_bullets = False  # Flag to skip markdown TOC list items

	skip_first_header = True  # Flag to skip the first header (document title)

	while idx < len(lines):

		line = lines[idx].rstrip()
		

		# Skip metadata section at the beginning

		if in_metadata_section:

			if line.startswith('**') and ':' in line:

				# Skip metadata lines

				idx += 1

				continue

			elif line.startswith('# '):

				# First header marks end of metadata

				in_metadata_section = False

				# Skip the first header (it's the document title)

				if skip_first_header:

					skip_first_header = False

					idx += 1

					continue

			elif line.strip() == '':

				# Skip empty lines in metadata section

				idx += 1

				continue

			else:

				# Non-header, non-metadata content

				if any(lines[j].startswith('**') and ':' in lines[j] for j in range(0, idx)):

					in_metadata_section = False
		

		# Skip empty lines

		if not line:

			idx += 1

			continue
		

		# Headers

		if line.startswith('# '):

			# H1 - Main title

			text = clean_markdown_formatting(line[2:])

			text = remove_header_numbering(text)

			para = doc.add_heading(text, level=1)
			

			# Add bookmark to heading

			bookmark_name = generate_bookmark_name(text)

			if bookmark_name:

				add_bookmark_to_paragraph(para, bookmark_name)
			

			# If this is "Obsah" (Contents), disable numbering and skip following list items

			if text.strip().lower() == 'obsah':

				# Disable numbering for this heading

				p_element = para._element

				pPr = p_element.find(qn('w:pPr'))

				if pPr is None:

					pPr = OxmlElement('w:pPr')

					p_element.insert(0, pPr)
				

				# Remove existing numPr if present

				numPr = pPr.find(qn('w:numPr'))

				if numPr is not None:

					pPr.remove(numPr)

				skip_toc_bullets = True

			else:

				skip_toc_bullets = False
			

		elif line.startswith('## '):

			# H2

			text = clean_markdown_formatting(line[3:])

			text = remove_header_numbering(text)

			para = doc.add_heading(text, level=2)
			

			# Add bookmark to heading

			bookmark_name = generate_bookmark_name(text)

			if bookmark_name:

				add_bookmark_to_paragraph(para, bookmark_name)
			

		elif line.startswith('### '):

			# H3

			text = clean_markdown_formatting(line[4:])

			text = remove_header_numbering(text)

			para = doc.add_heading(text, level=3)
			

			# Add bookmark to heading

			bookmark_name = generate_bookmark_name(text)

			if bookmark_name:

				add_bookmark_to_paragraph(para, bookmark_name)
			

		elif line.startswith('#### '):

			# H4

			text = clean_markdown_formatting(line[5:])

			text = remove_header_numbering(text)

			para = doc.add_heading(text, level=4)
			

			# Add bookmark to heading

			bookmark_name = generate_bookmark_name(text)

			if bookmark_name:

				add_bookmark_to_paragraph(para, bookmark_name)
		

		# Tables

		elif line.startswith('|'):

			table_data, new_idx = parse_table(lines, idx)

			if table_data:

				add_table_to_doc(doc, table_data)

				idx = new_idx - 1
		

		# Numbered Lists

		elif re.match(r'^\s*\d+\.\s+', line):

			# Skip TOC list items after Obsah header

			if skip_toc_bullets:

				idx += 1

				continue
			

			# Determine indentation level (0 spaces = level 1, 2 spaces = level 2, 4 spaces = level 3, etc.)

			leading_spaces = len(line) - len(line.lstrip())

			list_level = (leading_spaces // 2) + 1  # 0 spaces = level 1, 2 spaces = level 2, etc.
			

			# Extract text after number and dot

			list_text = re.sub(r'^\s*\d+\.\s+', '', line)
			

			# Check if this line contains internal links

			has_internal_links = re.search(r'\[([^\]]+)\]\(#([^\)]+)\)', list_text)
			

			para = doc.add_paragraph()
			

			if has_internal_links:

				# Process with link handling - first clean other markdown formatting

				text = re.sub(r'\*\*(.+?)\*\*', r'\1', list_text)  # Remove bold

				text = re.sub(r'\*(.+?)\*', r'\1', text)  # Remove italic

				text = re.sub(r'`(.+?)`', r'\1', text)  # Remove code

				# Process text with links

				process_text_with_links(para, text, heading_bookmarks)

			else:

				# No links - use regular markdown cleaning

				text = clean_markdown_formatting(list_text)

				para.add_run(text)
			

			# Try to apply List Number style first, otherwise fall back to manual numbering

			try:

				para.style = 'List Number' if list_level == 1 else f'List Number {list_level}'

			except KeyError:

				# Style not available, apply manual numbering with high num_id to avoid header conflicts

				apply_list_numbering(para, num_id=10, level=list_level - 1)
			

		# Bullet Lists

		elif line.startswith('- ') or (line.startswith('  ') and '- ' in line):

			# Skip TOC list items after Obsah header

			if skip_toc_bullets:

				idx += 1

				continue
			

			# Determine indentation level (0 spaces = level 1, 2 spaces = level 2, 4 spaces = level 3, etc.)

			leading_spaces = len(line) - len(line.lstrip())

			list_level = (leading_spaces // 2) + 1  # 0 spaces = level 1, 2 spaces = level 2, etc.
			

			list_text = line.lstrip()[2:]  # Remove leading spaces and '- '
			

			# Check if this line contains internal links

			has_internal_links = re.search(r'\[([^\]]+)\]\(#([^\)]+)\)', list_text)
			

			para = doc.add_paragraph()
			

			if has_internal_links:

				# Process with link handling - first clean other markdown formatting

				text = re.sub(r'\*\*(.+?)\*\*', r'\1', list_text)  # Remove bold

				text = re.sub(r'\*(.+?)\*', r'\1', text)  # Remove italic

				text = re.sub(r'`(.+?)`', r'\1', text)  # Remove code

				# Process text with links

				process_text_with_links(para, text, heading_bookmarks)

			else:

				# No links - use regular markdown cleaning

				text = clean_markdown_formatting(list_text)

				para.add_run(text)
			

			# Try to apply appropriate list style based on level

			if list_level == 1:

				try:

					para.style = 'List Bullet'

				except KeyError:

					# Style not available, add bullet manually

					para.paragraph_format.left_indent = Inches(0.25)

					# Prepend bullet to first run

					if para.runs:

						para.runs[0].text = '• ' + para.runs[0].text

			elif list_level == 2:

				try:

					para.style = 'List Bullet 2'

				except KeyError:

					# Style not available, add bullet manually with indent

					para.paragraph_format.left_indent = Inches(0.5)

					# Prepend bullet to first run

					if para.runs:

						para.runs[0].text = '◦ ' + para.runs[0].text

			else:

				# Level 3 or deeper

				try:

					para.style = f'List Bullet {list_level}'

				except KeyError:

					# Style not available, add bullet manually with deeper indent

					para.paragraph_format.left_indent = Inches(0.25 * list_level)

					# Prepend bullet to first run

					if para.runs:

						para.runs[0].text = '▪ ' + para.runs[0].text
			

		# Images

		elif line.startswith('<div align="center">'):

			# Look for img tag in following lines

			img_idx = idx + 1

			while img_idx < len(lines) and not lines[img_idx].endswith('</div>'):

				img_line = lines[img_idx].strip()

				if '<img' in img_line and 'src=' in img_line:

					# Extract image path from img tag

					match = re.search(r'src=["\'](.+?)["\']', img_line)

					if match:

						img_path = match.group(1)

						# Resolve path relative to markdown file

						md_dir = Path(md_file).parent

						full_img_path = md_dir / img_path

						

						if full_img_path.exists():

							try:

								# Add paragraph with center alignment

								para = doc.add_paragraph()

								para.alignment = WD_ALIGN_PARAGRAPH.CENTER

								run = para.add_run()

								# Add image at original size

								run.add_picture(str(full_img_path))

								print(f"Added image: {img_path}")

							except Exception as e:

								print(f"Warning: Could not add image {img_path}: {e}")

						else:

							print(f"Warning: Image file not found: {full_img_path}")

				

				img_idx += 1

			

			# Skip to end of div

			idx = img_idx
		

		# Regular paragraphs

		else:

			# Skip if it's a metadata line (shouldn't reach here, but just in case)

			if not in_metadata_section and line.startswith('**') and ':' in line:

				label, value = extract_metadata_value(line)

				if label in ['Projekt', 'Objednatel', 'Zhotovitel', 'Verze']:

					idx += 1

					continue
			

			# Check if this line contains internal links

			has_internal_links = re.search(r'\[([^\]]+)\]\(#([^\)]+)\)', line)
			

			# Check if this is a note (POZN or text in square brackets with italics)

			# Pattern: *POZN. ...* or *[text]*

			is_note = False

			if (line.strip().startswith('*POZN') and line.strip().endswith('*')) or (line.strip().startswith('*[') and line.strip().endswith(']*')):
				is_note = True            

			para = doc.add_paragraph()
			

			if has_internal_links:

				# Process with link handling - first clean other markdown formatting

				text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)  # Remove bold

				text = re.sub(r'\*(.+?)\*', r'\1', text)  # Remove italic

				text = re.sub(r'`(.+?)`', r'\1', text)  # Remove code

				# Strip square brackets from notes

				if is_note:

					text = text.replace('[', '').replace(']', '')

				# Process text with links

				process_text_with_links(para, text, heading_bookmarks)

			else:

				# No links - use regular markdown cleaning

				text = clean_markdown_formatting(line)

				# Strip square brackets from notes

				if is_note:

					text = text.replace('[', '').replace(']', '')

				if text:

					para.add_run(text)
			

			if is_note:

				try:

					para.style = 'Poznámka'

				except KeyError:

					# Style not available, make italic as fallback

					for run in para.runs:

						run.italic = True
		

		idx += 1
	

	# Save document

	print(f"Saving document to: {output_file}")

	os.makedirs(output_file.parent, exist_ok=True)

	doc.save(output_file)

	print("✓ Conversion complete!")



if __name__ == "__main__":

	# Check if files exist for PRE-BEP conversion
	if not INPUT_PREBEP_MD.exists():
		print(f"Error: Markdown file not found: {INPUT_PREBEP_MD}")
		exit(1)  
	if not TEMPLATE_PREBEP_DOCX.exists():
		print(f"Error: Template file not found: {TEMPLATE_PREBEP_DOCX}")
		exit(1)
	
	if not INPUT_EIR_MD.exists():
		print(f"Error: Markdown file not found: {INPUT_EIR_MD}")
		exit(1)
	if not TEMPLATE_EIR_DOCX.exists():
		print(f"Error: Template file not found: {TEMPLATE_EIR_DOCX}")
		exit(1)
	if not INPUT_PRILOHA_MD.exists():
		print(f"Error: Markdown file not found: {INPUT_PRILOHA_MD}")
		exit(1)
	if not TEMPLATE_PRILOHA_DOCX.exists():
		print(f"Error: Template file not found: {TEMPLATE_PRILOHA_DOCX}")
		exit(1)
	if not INPUT_PROTOKOL_MD.exists():
	    print(f"Error: Markdown file not found: {INPUT_PROTOKOL_MD}")
	    exit(1)
	if not TEMPLATE_PROTOKOL_DOCX.exists():
	    print(f"Error: Template file not found: {TEMPLATE_PROTOKOL_DOCX}")
	    exit(1)
	

	# Convert

	convert_markdown_to_word(INPUT_PREBEP_MD, TEMPLATE_PREBEP_DOCX, OUTPUT_PREBEP_DOCX)
	convert_markdown_to_word(INPUT_EIR_MD, TEMPLATE_EIR_DOCX, OUTPUT_EIR_DOCX)
	convert_markdown_to_word(INPUT_PROTOKOL_MD, TEMPLATE_PROTOKOL_DOCX, OUTPUT_PROTOKOL_DOCX)
	convert_markdown_to_word(INPUT_PRILOHA_MD, TEMPLATE_PRILOHA_DOCX, OUTPUT_PRILOHA_DOCX)
